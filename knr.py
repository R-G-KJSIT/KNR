import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
from PIL import Image
from groq import Groq
import re

# Constants
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]  # Replace with your actual API key

# Initialize the Groq client
client = Groq(api_key=GROQ_API_KEY)

# Azure AD details for Microsoft Graph API
client_id = st.secrets["client_id"]
client_secret = st.secrets['client_secret']
scopes = ['wl.signin', 'wl.offline_access', 'onedrive.readwrite']
tenant_id = st.secrets["tenant_id"]
scope = "https://graph.microsoft.com/.default"

def get_access_token():
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    response = requests.post(token_url, data={
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': scope
    })
    response.raise_for_status()  # Raise an exception for HTTP errors
    return response.json().get('access_token')

# Placeholder for Microsoft Graph API
GRAPH_API_BASE_URL = "https://graph.microsoft.com/v1.0/me/drive/root:/Documents/College Documents/S4DS/Knowhow Reports/"

# Function to call Groq API using the Groq module
def call_groq_api(points):
    prompt = (
        "I am about to give you a few points that summarize an event. "
        "Please format these points into a more professional style. "
        "Remember this is for a report so whatever points are given, make sure you write it in a way that displays what was taught/done and not what the points mean (Like write something along the lines of '_____ was conducted where ________ was taught and ____ was performed')"
        "Provide them in a bullet-point format, separated by ';'. "
        "Make sure the words you use are extremely simple but professional. Don't use any complex words and don't make 1 point use longer than 10 words (Make sure it STRICTLY ENDS BEFORE 10 Words. Not a singlge word more)"
        "If any text should be bold, enclose it in '**'. "
        "Output nothing other than the refined points in the specified format. And I mean ABSOLUTELY NOTHING, not even things like 'Here's the .....'"
        "Make sure you make a long sentence of atleast 10 words per point"
        "Points: "
    )
    messages = [
        {
            "role": "user",
            "content": f"{prompt} {', '.join(points)}"
        }
    ]
    
    try:
        chat_completion = client.chat.completions.create(
            messages=messages,
            model="llama3-70b-8192"
        )
        response = chat_completion.choices[0].message.content.strip()
        return response
    except Exception as e:
        st.error(f"Error calling Groq API: {e}")
        return ""

# Function to create DOCX with title, conducted_by, bullet points, and venue
def create_docx(title, conducted_by, points, venue):
    doc = Document()
    
    # Add logos to header
    header = doc.sections[0].header
    
    # Helper function to add images to the header
    def add_image_to_header(image_path, width=Pt(50)):
        try:
            with Image.open(image_path) as img:
                image_stream = BytesIO()
                img.save(image_stream, format='PNG' if image_path.endswith('.png') else 'JPEG')
                header_paragraph = header.add_paragraph()
                header_paragraph.add_run().add_picture(image_stream, width=width)
        except Exception as e:
            st.error(f"Error adding image to document: {e}")

    # Add the first logo
    add_image_to_header("collegelogo.png", width=Pt(50))
    
    # Add space between the images
    header.add_paragraph()  # Add a blank paragraph to create space
    
    # Add the second logo
    add_image_to_header("logo.jpg", width=Pt(50))

    # Apply default font settings
    def set_font(run):
        run.font.name = 'Segoe UI'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Segoe UI')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Segoe UI')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Segoe UI')

    # Add title
    heading = doc.add_heading(title, level=1)
    heading.alignment = 1  # Center alignment
    for run in heading.runs:
        set_font(run)
    
    # Add date and venue
    date_venue_paragraph = doc.add_paragraph()
    date_venue_paragraph.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
    date_venue_paragraph.add_run(f"Venue: {venue}\n")
    date_venue_paragraph.alignment = 0  # Left alignment
    for run in date_venue_paragraph.runs:
        set_font(run)

    # Add conducted by line
    conducted_by_paragraph = doc.add_paragraph(f"Conducted by: {conducted_by}", style='Normal')
    for run in conducted_by_paragraph.runs:
        set_font(run)

    # Parse and add bullet points
    if points:
        doc.add_paragraph("Event Highlights:", style='Normal')
        points_list = points.split(';')
        for point in points_list:
            point = point.strip()
            if point:
                p = doc.add_paragraph(style='ListBullet')
                runs = re.split(r'(\*\*.*?\*\*)', point)  # Split based on bold markers
                for run in runs:
                    if run.startswith('**') and run.endswith('**'):
                        bold_run = p.add_run(run[2:-2])
                        bold_run.bold = True
                        set_font(bold_run)
                    else:
                        regular_run = p.add_run(run)
                        set_font(regular_run)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

    doc = Document()
    
    # Add logo to header
    header = doc.sections[0].header
    logo_path = "logo.jpg"  # Ensure this path is correct
    try:
        logo = Image.open(logo_path)
        logo_stream = BytesIO()
        logo.save(logo_stream, format='JPEG')
        header_paragraph = header.add_paragraph()
        header_paragraph.add_run().add_picture(logo_stream, width=Pt(50))  # Adjust size as needed
    except Exception as e:
        st.error(f"Error adding logo to document: {e}")

    # Add title
    heading = doc.add_heading(title, level=1)
    heading.alignment = 1  # Center alignment
    run = heading.runs[0]
    run.font.name = 'Segoe UI'
    run.font.size = Pt(14)
    
    # Add date and venue
    date_venue_paragraph = doc.add_paragraph()
    date_venue_paragraph.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
    date_venue_paragraph.add_run(f"Venue: {venue}\n")
    date_venue_paragraph.alignment = 0  # Left alignment

    # Add conducted by line
    doc.add_paragraph(f"Conducted by: {conducted_by}", style='Normal')

    # Parse and add bullet points
    if points:
        doc.add_paragraph("Event Highlights:", style='Normal')
        points_list = points.split(';')
        for point in points_list:
            point = point.strip()
            if point:
                p = doc.add_paragraph(style='ListBullet')
                runs = re.split(r'(\*\*.*?\*\*)', point)  # Split based on bold markers
                for run in runs:
                    if run.startswith('**') and run.endswith('**'):
                        p.add_run(run[2:-2]).bold = True
                    else:
                        p.add_run(run)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to upload DOCX to OneDrive
def upload_to_onedrive(file_buffer, filename):
    access_token = get_access_token()
    upload_url = f"{GRAPH_API_BASE_URL}{filename}:/content"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/octet-stream"
    }
    response = requests.put(upload_url, headers=headers, data=file_buffer)
    if response.status_code == 201:
        st.success("File uploaded successfully!")
    else:
        st.error(f"Error uploading file to OneDrive: {response.status_code} - {response.text}")

# Streamlit UI
st.title("Knowhow Workshop Reporting Software")

# Use session state to persist points
if 'points' not in st.session_state:
    st.session_state.points = []

# Date picker
event_date = st.date_input("Date of Event")

# Domain Selector
domain = st.selectbox(
    "Domain",
    ["AI-ML", "DataScience and Analytics", "Cloud Native", "Web Development", "AR/VR", "IoT", "Blockchain", "Cybersecurity"]
)

# Conducted By Input
conducted_by = st.text_input("Conducted By")

# Venue Input
venue = st.text_input("Venue")

# Format the date and domain for the file name
formatted_date = event_date.strftime("%Y-%m-%d")
domain_slug = domain.replace(" ", "_")  # Replace spaces with underscores
filename = f"report_{formatted_date}_{domain_slug}.docx"
title = f"Knowhow {domain} workshop {formatted_date}"  # Title format

# Brief Points Input
st.subheader("Brief Points")
point_input = st.text_input("Enter a point", "")

if st.button("Add Point"):
    if point_input:
        st.session_state.points.append(point_input)
        st.rerun()  # Refresh the app to display updated points
    else:
        st.warning("Please enter a point before adding.")

# Display points
if st.session_state.points:
    for i, point in enumerate(st.session_state.points):
        st.text(f"Point {i + 1}: {point}")

# Generate Report Button
if st.button("Generate Report"):
    if len(st.session_state.points) == 0:
        st.warning("Please add at least one point.")
    elif not conducted_by:
        st.warning("Please enter the name of the person who conducted the event.")
    elif not venue:
        st.warning("Please enter the venue.")
    else:
        # Call Groq API to format the points
        raw_points = call_groq_api(st.session_state.points)
        
        # Ensure at least 5 points
        points_list = raw_points.split(';')
        if len(points_list) < 5:
            additional_points = ["Additional point placeholder." for _ in range(5 - len(points_list))]
            points_list.extend(additional_points)
        
        # Create DOCX with title, conducted_by, formatted points, and venue
        docx_buffer = create_docx(title, conducted_by, ';'.join(points_list), venue)
        
        # Create a download button
        st.download_button(
            label="Download Report",
            data=docx_buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Upload to OneDrive with dynamic filename
        # Since upload_to_onedrive needs a file buffer, we need to reset buffer's position
        docx_buffer.seek(0)
        upload_to_onedrive(docx_buffer, filename)
